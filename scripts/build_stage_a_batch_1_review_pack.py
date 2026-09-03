import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / 'backend/src/modules/nutrition/food-curation/data/stage-a.batch-1.pending-approval.json'
OUT = ROOT / 'backend/src/modules/nutrition/FITEATSY_STAGE_A_NUTRITIONIST_APPROVAL_PACK_BATCH_1_v1.docx'
TASK = ROOT / 'backend/src/modules/nutrition/FITEATSY_CHAPATI_PROTOCOL_COMPLIANT_REMEASUREMENT_TASK_v1.docx'
NAVY = RGBColor(11, 37, 69); BLUE = RGBColor(46, 116, 181); MUTED = RGBColor(90, 99, 112); GOLD = RGBColor(122, 90, 0)

def set_font(run, size=10, bold=False, color=None):
    run.font.name = 'Calibri'; run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Calibri'); run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size); run.bold = bold
    if color: run.font.color.rgb = color

def setup(doc, label):
    sec = doc.sections[0]; sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(.492)
    normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(10); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after, color in [('Heading 1',16,18,10,BLUE),('Heading 2',13,14,7,BLUE),('Heading 3',12,10,5,NAVY)]:
        style = doc.styles[name]; style.font.name = 'Calibri'; style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = color; style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after); style.paragraph_format.keep_with_next = True
    h = sec.header.paragraphs[0]; set_font(h.add_run('FITEATSY  |  CONTROLLED FOOD CURATION'), 8, True, MUTED)
    f = sec.footer.paragraphs[0]; f.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_font(f.add_run(label), 8, False, MUTED)

def title(doc, kicker, heading, subtitle):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(24); p.paragraph_format.space_after = Pt(4); set_font(p.add_run(kicker), 10, True, BLUE)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5); set_font(p.add_run(heading), 24, True, NAVY)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(16); set_font(p.add_run(subtitle), 13, False, BLUE)

def table(doc, rows, widths=(2.0,4.5)):
    t = doc.add_table(rows=0, cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.autofit = False
    for label, value in rows:
        cells = t.add_row().cells
        for cell, width in zip(cells, widths): cell.width = Inches(width); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(cells[0].paragraphs[0].add_run(str(label)), 9, True, NAVY)
        set_font(cells[1].paragraphs[0].add_run(str(value)), 9, False, MUTED)
        for cell in cells:
            tc = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'F8FAFC'); tc.append(shd)
    if t.rows:
        header = t.rows[0]._tr.get_or_add_trPr(); repeat = OxmlElement('w:tblHeader'); repeat.set(qn('w:val'), 'true'); header.append(repeat)
    pr=t._tbl.tblPr; w=pr.first_child_found_in('w:tblW'); w.set(qn('w:type'),'dxa'); w.set(qn('w:w'),'9360')
    ind=OxmlElement('w:tblInd'); ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa'); pr.append(ind)
    grid=t._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for dxa in (2880,6480): col=OxmlElement('w:gridCol'); col.set(qn('w:w'),str(dxa)); grid.append(col)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def checkbox(doc, text):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); set_font(p.add_run('☐  '+text),10)

def page(doc, heading, subtitle):
    doc.add_page_break(); doc.add_heading(heading, level=1)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(9); set_font(p.add_run(subtitle),10,False,MUTED)

def build_pack():
    pack=json.loads(DATA.read_text())
    doc=Document(); setup(doc,'Stage A Nutritionist Approval Pack — Batch 1')
    title(doc,'STAGE A REVIEW','Batch 1 Formula Approval Pack','Five controlled reference Foods | Decision required')
    table(doc,[('Pack ID',pack['packId']),('Status','PENDING — NO DECISION PRESELECTED'),('Decision options','APPROVED / CHANGES_REQUIRED / REJECTED'),('Formula hash','Generated only after qualified approval'),('Production effect','NONE')])
    p=doc.add_paragraph(); set_font(p.add_run('Reviewer instruction: assess each formula independently. Approval binds only the exact material formula; measurements, calculation, Stage B review, and release remain separate gates.'),10,True,GOLD)
    page(doc,'Reviewer Identity & Declaration','Complete once for the pack; record an individual decision for every Food.')
    table(doc,[('Reviewer name / ID','____________________________________________'),('Qualification','____________________________________________'),('Qualification reference','____________________________________________'),('Review date','YYYY-MM-DD: ______________________________'),('Declaration','I reviewed each exact formula and recorded my independent decision.'),('Signature / recorded approval reference','____________________________________________')])
    for formula in pack['formulas']:
        page(doc,formula['name'],f"{formula['preparationId']} | Proposed formula {formula['formulaVersion']}")
        ingredients='; '.join(f"{i['identity']} — {i['quantityGrams']} g ({i['role']})" for i in formula['ingredients'])
        water=f"{formula['water']['quantityMl']} ml — {formula['water']['handling']}"
        rows=[('Preparation ID',formula['preparationId']),('Formula Version',formula['formulaVersion']),('Exact proposed ingredients',ingredients),('Preparation method',formula['preparationMethod']),('Water handling',water),('Serving concept',formula['servingConcept']),('Present components',', '.join(formula['hardContext']['present']) or 'None declared'),('Absent verified components',', '.join(formula['hardContext']['absentVerified']) or 'None declared'),('Allergens',', '.join(formula['hardContext']['allergens']) or 'None declared'),('Source dependencies',', '.join(formula['sourceDependencies'])),('Known limitations',', '.join(formula['limitations']) or 'None recorded')]
        if 'postCookingFat' in formula: rows.insert(5,('Post-cooking fat',f"Applied: YES | Quantity: {'MISSING' if formula['postCookingFat']['quantityGrams'] is None else str(formula['postCookingFat']['quantityGrams'])+' g'}"))
        table(doc,rows)
        doc.add_heading('Decision',level=2)
        for choice in pack['decisionOptions']: checkbox(doc,choice)
        table(doc,[('Reason','____________________________________________\n____________________________________________'),('Reviewer initials / reference','____________________________________________'),('Date','YYYY-MM-DD: ______________________________')])
    doc.core_properties.title='FITEATSY Stage A Nutritionist Approval Pack Batch 1 v1'; doc.core_properties.author='FITEATSY'
    doc.save(OUT)

def build_task():
    doc=Document(); setup(doc,'Chapati Protocol-Compliant Remeasurement Task v1')
    title(doc,'HUMAN TASK','Chapati Remeasurement','CP_CHAPATI | Repeat Chapati only')
    p=doc.add_paragraph(); set_font(p.add_run('Do not repeat Moong Dal, Bhindi Sabji, Bhindi Aloo, or Peanut Poha.'),10,True,GOLD)
    table(doc,[('Current finding','Four measured pieces reconcile to 140 g, but protocol requires at least five independently formed pieces.'),('Formula prerequisite','Use the exact Stage A-approved Formula Version and Formula SHA-256.'),('Post-cooking fat','Measure quantitatively if applied. Record 0 g only if physically verified and formula-consistent.'),('Missing audit metadata','Operator, date, equipment ID, scale resolution.')])
    doc.add_heading('Repeat evidence',level=1)
    table(doc,[('Formula Version / SHA-256','____________________________________________'),('Actual atta','____________ g'),('Actual water','____________ ml'),('Actual kneading oil/ghee','____________ g'),('Actual dusting flour','____________ g'),('Post-cooking fat','____________ g | physically verified 0 g if none'),('Final cooked batch','____________ g'),('Produced pieces','____________ (minimum 5)'),('Piece observations','1 ____ g   2 ____ g   3 ____ g   4 ____ g   5 ____ g\nAdditional: __________________________________'),('Operator','____________________________________________'),('Measurement date','YYYY-MM-DD: ______________________________'),('Equipment ID','____________________________________________'),('Scale resolution','____________ g'),('Deviations','____________________________________________')])
    doc.add_heading('Operator declaration',level=1)
    for text in ['Every produced piece listed was independently weighed.','At least five independently formed pieces were measured.','Post-cooking fat quantity is measured, not inferred.','Final batch yield is physically measured.','No missing value was copied, averaged, or invented.']: checkbox(doc,text)
    table(doc,[('Signature / recorded declaration','____________________________________________'),('Date','YYYY-MM-DD: ______________________________')])
    doc.core_properties.title='FITEATSY Chapati Protocol-Compliant Remeasurement Task v1'; doc.core_properties.author='FITEATSY'
    doc.save(TASK)

if __name__=='__main__':
    build_pack(); build_task(); print(OUT); print(TASK)
