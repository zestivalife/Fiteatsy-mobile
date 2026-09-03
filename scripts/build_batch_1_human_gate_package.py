from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "backend/src/modules/nutrition/FITEATSY_BATCH_1_HUMAN_GATE_PACKAGE_v1.docx"
BLUE = "2E74B5"
NAVY = "0B2545"
PALE = "E8EEF5"
GRAY = "666666"
WIDTHS = [2450, 6910]


def set_font(run, size=10.5, bold=False, color="000000"):
    run.font.name = "Calibri"
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None:
        tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")
    if tc_w.getparent() is None:
        tc_pr.append(tc_w)


def set_table_geometry(table):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    if tbl_w.getparent() is None:
        tbl_pr.append(tbl_w)
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    if tbl_ind.getparent() is None:
        tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in WIDTHS:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, WIDTHS[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)


def add_row(table, label, value, label_fill=PALE):
    left, right = table.add_row().cells
    shade(left, label_fill)
    set_font(left.paragraphs[0].add_run(label), bold=True, color=BLUE)
    set_font(right.paragraphs[0].add_run(value))


def add_decision_card(doc, title, rows):
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(5)
    set_font(heading.add_run(title), 12, True, NAVY)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        add_row(table, label, value)
    set_table_geometry(table)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(section, side, Inches(0.8))
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.15
for name, size, before, after in (("Heading 1", 16, 14, 7), ("Heading 2", 13, 10, 5)):
    style = doc.styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(BLUE)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("FITEATSY | Controlled Food Curation"), 9, color=GRAY)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("Candidate-only human decision package | No production activation"), 8.5, color=GRAY)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(3)
set_font(title.add_run("BATCH 1 HUMAN GATE PACKAGE"), 21, True, NAVY)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(11)
set_font(subtitle.add_run("Stage A formula decisions + source identity decisions"), 11, True, BLUE)

doc.add_heading("A. Stage A Food Decisions", level=1)
reviewer = doc.add_table(rows=0, cols=2)
reviewer.style = "Table Grid"
for label, value in (
    ("Reviewer name / ID", "____________________________________________"),
    ("Qualification + reference", "Qualification: ____________________   Reference: ____________________"),
    ("Review date + declaration", "Date (YYYY-MM-DD): ______________   Declaration / signature: ____________________"),
):
    add_row(reviewer, label, value)
set_table_geometry(reviewer)

stage_a = [
    ("Chapati", "CP_CHAPATI_STAGE_A_v1"),
    ("Moong Dal", "CP_MOONG_DAL_STAGE_A_v1"),
    ("Bhindi Sabji", "CP_BHINDI_SABJI_STAGE_A_v1"),
    ("Bhindi Aloo", "CP_BHINDI_ALOO_STAGE_A_v1"),
    ("Peanut Poha", "CP_POHA_PEANUT_STAGE_A_v1"),
]
for food, version in stage_a:
    add_decision_card(doc, food, [
        ("Formula version", version),
        ("Decision", "[  ] APPROVED     [  ] CHANGES_REQUIRED     [  ] REJECTED"),
        ("Reason + reviewer record", "Reason / changes: ______________________________   Initials + date: ____________________"),
    ])

doc.add_page_break()
doc.add_heading("B. Source Identity Decisions", level=1)
source_reviewer = doc.add_table(rows=0, cols=2)
source_reviewer.style = "Table Grid"
for label, value in (
    ("Reviewer name / ID", "Priyanshi Srivastava"),
    ("Qualification + reference", "B. Pharma and MSc DFSM | 2015 and 2024"),
    ("Review date + declaration", "2026-09-03 | Priyanshi Srivastava"),
):
    add_row(source_reviewer, label, value)
set_table_geometry(source_reviewer)

sources = [
    ("Refined Sunflower Oil", "USDA FDC 1750349 | Oil, sunflower | core incomplete", "[  ] APPROVE EXACT MAPPING   [  ] REJECT MAPPING   [  ] REQUEST ALTERNATE SOURCE"),
    ("Cow Ghee", "USDA FDC 171314 | Butter, Clarified butter (ghee) | Cow species unresolved", "[  ] APPROVE EXACT MAPPING   [  ] REJECT MAPPING   [  ] REQUEST ALTERNATE SOURCE"),
    ("Groundnut Oil", "USDA FDC 171410 | Oil, peanut, salad or cooking | grade/process unresolved", "[  ] APPROVE EXACT MAPPING   [  ] REJECT MAPPING   [  ] REQUEST ALTERNATE SOURCE"),
    ("Split Hulled Yellow Moong Dal", "No acceptable approved exact generic source", "[  ] CONFIRM NO MATCH   [  ] PROVIDE APPROVED EXACT SOURCE"),
    ("Dry Flattened Rice / Poha", "No acceptable approved exact generic source", "[  ] CONFIRM NO MATCH   [  ] PROVIDE APPROVED EXACT SOURCE"),
]
for ingredient, candidate, choices in sources:
    add_decision_card(doc, ingredient, [
        ("Current candidate", candidate),
        ("Decision", choices),
        ("Alternate exact source evidence", "Source / record / version / exact description / rights / nutrient basis: __________________________"),
        ("Reviewer initials + date", "____________________________________________"),
    ])

doc.core_properties.title = "FITEATSY Batch 1 Human Gate Package v1"
doc.core_properties.author = "FITEATSY"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
