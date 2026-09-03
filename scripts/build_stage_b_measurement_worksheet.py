from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parents[1] / 'backend/src/modules/nutrition/FITEATSY_CONTROLLED_REFERENCE_BATCH_1_PHYSICAL_MEASUREMENT_WORKSHEET.docx'
NAVY = RGBColor(11, 37, 69); BLUE = RGBColor(46, 116, 181); MUTED = RGBColor(90, 99, 112)

def font(run, size=10, bold=False, color=None):
    run.font.name = 'Calibri'; run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Calibri'); run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size); run.bold = bold
    if color: run.font.color.rgb = color

def field(doc, label, value='________________________________________________________'):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    font(p.add_run(label + ': '), bold=True, color=NAVY); font(p.add_run(value), color=MUTED)

def heading(doc, text):
    p = doc.add_paragraph(style='Heading 1'); p.paragraph_format.keep_with_next = True
    font(p.add_run(text), size=16, bold=True, color=BLUE)

def table(doc, rows):
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.autofit = False
    t.columns[0].width = Inches(2.15); t.columns[1].width = Inches(4.35)
    for label, value in rows:
        cells = t.add_row().cells; cells[0].width = Inches(2.15); cells[1].width = Inches(4.35)
        cells[0].vertical_alignment = cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        font(cells[0].paragraphs[0].add_run(label), bold=True, color=NAVY)
        font(cells[1].paragraphs[0].add_run(value), color=MUTED)
    t._tbl.remove(t.rows[0]._tr)
    first_row_properties = t.rows[0]._tr.get_or_add_trPr(); first_row_properties.append(OxmlElement('w:tblHeader'))
    tbl_pr = t._tbl.tblPr; width = tbl_pr.first_child_found_in('w:tblW'); width.set(qn('w:type'), 'dxa'); width.set(qn('w:w'), '9360')
    ind = OxmlElement('w:tblInd'); ind.set(qn('w:w'), '120'); ind.set(qn('w:type'), 'dxa'); tbl_pr.append(ind)
    for row in t.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'F8FAFC'); tc_pr.append(shd)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def checkbox(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4); font(p.add_run('☐  ' + text), size=10)

def new_page(doc, title, subtitle):
    doc.add_page_break(); heading(doc, title)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10); font(p.add_run(subtitle), color=MUTED)

def build():
    doc = Document(); sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(.492)
    normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(10); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
    h = sec.header.paragraphs[0]; font(h.add_run('FITEATSY  |  CONTROLLED FOOD MEASUREMENT'), size=8, bold=True, color=MUTED)
    f = sec.footer.paragraphs[0]; f.alignment = WD_ALIGN_PARAGRAPH.RIGHT; font(f.add_run('Batch 1 Physical Measurement Worksheet'), size=8, color=MUTED)

    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(28); font(p.add_run('MEASUREMENT INPUT'), size=11, bold=True, color=BLUE)
    p = doc.add_paragraph(); font(p.add_run('Controlled Reference Batch 1'), size=26, bold=True, color=NAVY)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14); font(p.add_run('Physical Measurement Worksheet'), size=15, color=BLUE)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(12); font(p.add_run('NOT CANONICAL EVIDENCE UNTIL COMPLETED AND VALIDATED'), size=11, bold=True, color=RGBColor(122,90,0))
    table(doc, [('Batch ID', 'FITEATSY-FIRST-FIVE-v1'), ('Formula Version references', 'Record exact approved versions below'), ('Measurement methodology', 'FITEATSY Controlled Food Measurement Protocol v1'), ('Operator', '________________________________'), ('Date', 'YYYY-MM-DD: __________________'), ('Scale / Equipment ID', '________________________________'), ('Scale resolution', '____________ g'), ('Reference Katori', 'FITEATSY_KATORI_V1')])
    p = doc.add_paragraph(); font(p.add_run('Do not enter Calories or macros. Do not copy proposed quantities as actual values unless physically measured.'), bold=True, color=NAVY)

    new_page(doc, 'Reference Katori', 'Calibrate the physical vessel before Food serving runs.')
    table(doc, [('Reference ID', 'FITEATSY_KATORI_V1'), ('Empty weight', '____________ g'), ('Capacity run 1', '____________ ml'), ('Capacity run 2', '____________ ml'), ('Capacity run 3', '____________ ml'), ('Selected capacity', 'SYSTEM / REVIEW RESULT'), ('Notes', '____________________________________________')])
    p = doc.add_paragraph(); font(p.add_run('The operator does not calculate a universal gram serving from vessel capacity.'), bold=True, color=NAVY)

    new_page(doc, 'Chapati', 'Record the approved formula reference and every physically measured value.')
    table(doc, [('Preparation ID', 'CP_CHAPATI'), ('Formula version', '________________'), ('Formula SHA-256', '________________________________________________________'), ('Measurement run ID', '________________'), ('Actual atta', '____________ g'), ('Actual water', '____________ ml'), ('Actual oil/ghee', '____________ g'), ('Dusting flour', '____________ g / NOT USED'), ('Final cooked batch', '____________ g'), ('Cooked pieces', '____________'), ('Piece weights', '1 ____ g   2 ____ g   3 ____ g   4 ____ g   5 ____ g'), ('Ghee/butter applied', 'YES / NO'), ('Formula deviation', 'YES / NO'), ('Notes', '____________________________________________')])

    new_page(doc, 'Moong Dal', 'Use the exact approved ingredient Food Version IDs; record actual quantities independently.')
    table(doc, [('Preparation ID', 'CP_MOONG_DAL'), ('Formula version', '________________'), ('Formula SHA-256', '________________________________________________________'), ('Measurement run ID', '________________'), ('Ingredient IDs + actual g', '____________________________________________'), ('Actual water', '____________ ml'), ('Actual oil', '____________ g'), ('Onion', '____________ g'), ('Garlic', '____________ g'), ('Final prepared weight', '____________ g'), ('Katori runs', '1 ____ g   2 ____ g   3 ____ g'), ('Formula deviation', 'YES / NO'), ('Notes', '____________________________________________')])

    new_page(doc, 'Bhindi Sabji', 'Potato must be a physically verified zero for the approved potato-free formula.')
    table(doc, [('Preparation ID', 'CP_BHINDI_SABJI'), ('Formula version / SHA-256', '____________________________________________'), ('Measurement run ID', '________________'), ('Ingredient IDs + actual g', '____________________________________________'), ('Actual water / oil', '________ ml  /  ________ g'), ('Potato actual', '____________ g  | expected controlled formula: 0 g'), ('Garlic actual', '____________ g'), ('Final prepared weight', '____________ g'), ('Katori runs', '1 ____ g   2 ____ g   3 ____ g'), ('Formula deviation', 'YES / NO'), ('Notes', '____________________________________________')])

    new_page(doc, 'Bhindi Aloo', 'Potato presence must be demonstrated by actual measured quantity.')
    table(doc, [('Preparation ID', 'CP_BHINDI_ALOO'), ('Formula version / SHA-256', '____________________________________________'), ('Measurement run ID', '________________'), ('Ingredient IDs + actual g', '____________________________________________'), ('Actual water / oil', '________ ml  /  ________ g'), ('Potato actual', '____________ g'), ('Garlic actual', '____________ g'), ('Final prepared weight', '____________ g'), ('Katori runs', '1 ____ g   2 ____ g   3 ____ g'), ('Formula deviation', 'YES / NO'), ('Notes', '____________________________________________')])

    new_page(doc, 'Peanut Poha', 'Peanut and dry Poha quantities must be measured; source dependencies remain independently governed.')
    table(doc, [('Preparation ID', 'CP_POHA_PEANUT'), ('Formula version / SHA-256', '____________________________________________'), ('Measurement run ID', '________________'), ('Ingredient IDs + actual g', '____________________________________________'), ('Dry Poha', '____________ g'), ('Peanut actual', '____________ g'), ('Actual water / oil', '________ ml  /  ________ g'), ('Potato / Garlic', '________ g  /  ________ g'), ('Final prepared weight', '____________ g'), ('Reference vessel', 'KATORI / BOWL / APPROVED OTHER: ____________'), ('Serving runs', '1 ____ g   2 ____ g   3 ____ g'), ('Formula deviation', 'YES / NO'), ('Notes', '____________________________________________')])

    new_page(doc, 'Completion Declaration', 'Complete only after all physical work and repeat measurements are finished.')
    for text in ['Values are physically measured.', 'Proposed formula values were not copied as actual values unless actually measured.', 'Final yield was physically weighed.', 'Serving measurements were physically repeated.', 'No Calories or macros were estimated.', 'All deviations were recorded.', 'No missing measurement was guessed.']: checkbox(doc, text)
    field(doc, 'Operator'); field(doc, 'Date', 'YYYY-MM-DD: ____________________________________________'); field(doc, 'Signature / recorded declaration')
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); font(p.add_run('NEXT: Validate the structured measurement input. Stage B calculation remains blocked until complete evidence passes validation.'), bold=True, color=NAVY)
    doc.core_properties.title = 'FITEATSY Controlled Reference Batch 1 Physical Measurement Worksheet'
    doc.core_properties.author = 'FITEATSY'
    doc.save(OUT); print(OUT)

if __name__ == '__main__': build()
