from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "backend/src/modules/nutrition/FITEATSY_BATCH_1_SOURCE_IDENTITY_REVIEW_TASK_v1_UPDATED.docx"
BLUE = "2E74B5"
PALE = "E8EEF5"
GRAY = "666666"
WIDTHS = [2300, 7060]

ITEMS = [
    ("REFINED_SUNFLOWER_OIL", "USDA FDC 1750349", "Oil, sunflower", "Refined grade is not established by the source description. The candidate also does not expose the complete mandatory Fiteatsy Energy/Protein/Carbohydrate/Fat/Fibre core vector. Even if identity is approved, calculation readiness remains blocked until the mandatory core vector is satisfied through an approved source path.", "CC0 approved", "Incomplete", "APPROVE EXACT MAPPING   /   REJECT MAPPING   /   REQUEST ALTERNATE SOURCE"),
    ("COW_GHEE", "USDA FDC 171314 (SR Legacy 2018-04)", "Butter, Clarified butter (ghee)", "This is a stronger preparation-name candidate than prior FDC 173412 (Butter oil, anhydrous), which remains in candidate history. The mandatory core vector is complete, but Cow-species identity is not established by the description. Human identity review remains required.", "CC0 approved", "Complete", "APPROVE EXACT MAPPING   /   REJECT MAPPING   /   REQUEST ALTERNATE SOURCE"),
    ("GROUNDNUT_OIL", "USDA FDC 171410", "Oil, peanut, salad or cooking", "Groundnut/peanut common-name equivalence is plausible, but exact grade/process is unspecified. Approval must confirm the candidate is sufficiently exact for the governed GROUNDNUT_OIL identity; otherwise request an alternate approved source.", "CC0 approved", "Complete", "APPROVE EXACT MAPPING   /   REJECT MAPPING   /   REQUEST ALTERNATE SOURCE"),
    ("SPLIT_HULLED_YELLOW_MOONG_DAL", "No exact approved generic record", "FDC 174256 is whole mature mung seed, raw", "Required preparation identity is split hulled yellow moong dal. Whole mature mung, whole green mung, generic lentils, and branded-product records are not acceptable generic canonical proxies.", "No adopted record", "Unavailable", "CONFIRM NO MATCH   /   PROVIDE APPROVED EXACT SOURCE"),
    ("DRY_FLATTENED_RICE_POHA", "No exact approved record", "No generic dry flattened-rice/Poha FDC record", "Required preparation identity is dry flattened rice / Poha. Generic rice, cooked rice, puffed rice, rice flour, noodles, and generic rice cereal are not acceptable proxies.", "No adopted record", "Unavailable", "CONFIRM NO MATCH   /   PROVIDE APPROVED EXACT SOURCE"),
]

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None:
        tc_pr.append(shd)

def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None: tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(width)); tc_w.set(qn("w:type"), "dxa")
    if tc_w.getparent() is None: tc_pr.append(tc_w)

def set_table_geometry(table):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None: tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), "9360"); tbl_w.set(qn("w:type"), "dxa")
    if tbl_w.getparent() is None: tbl_pr.append(tbl_w)
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None: tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120"); tbl_ind.set(qn("w:type"), "dxa")
    if tbl_ind.getparent() is None: tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in WIDTHS:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, WIDTHS[idx]); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_font(run, size=10.5, bold=False, color="000000"):
    run.font.name = "Calibri"; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri"); run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = RGBColor.from_string(color)

def add_kv(table, label, value, fill=None):
    cells = table.add_row().cells
    if fill: shade(cells[0], fill)
    p = cells[0].paragraphs[0]; p.paragraph_format.space_after = Pt(0); set_font(p.add_run(label), bold=True, color=BLUE)
    p = cells[1].paragraphs[0]; p.paragraph_format.space_after = Pt(0); set_font(p.add_run(value))

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"): setattr(sec, side, Inches(1))
sec.header_distance = Inches(.492); sec.footer_distance = Inches(.492)
styles = doc.styles
normal = styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
for name, size, before, after in (("Heading 1",16,18,10),("Heading 2",13,14,7)):
    st=styles[name]; st.font.name="Calibri"; st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(BLUE); st.font.bold=True; st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)

header = sec.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("FITEATSY | Controlled Food Curation"), 9, False, GRAY)
footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("Candidate-only review task | No production activation"), 8.5, False, GRAY)

p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); set_font(p.add_run("BATCH 1 SOURCE IDENTITY REVIEW TASK"), 22, True, "0B2545")
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(14); set_font(p.add_run("FITEATSY_SOURCE_IDENTITY_REVIEW_TASK_v1 | Status: PENDING HUMAN IDENTITY REVIEW"), 11, True, BLUE)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(10); set_font(p.add_run("Purpose. "), 11, True); set_font(p.add_run("Record only the unresolved source-identity decisions needed by the five-food methodology cohort. Rights approval does not establish Food identity. No choice is preselected. Current engineering state: 7 ingredient mappings READY, Water METHOD READY, 3 candidate mappings require identity review, and 2 exact source matches remain unavailable."), 11)

doc.add_heading("Reviewer record", level=1)
t=doc.add_table(rows=0, cols=2); t.style="Table Grid"
for label, value in (
    ("Reviewer name / ID", "Priyanshi Srivastava"),
    ("Qualification", "B. Pharma and MSc DFSM"),
    ("Qualification reference", "2015 and 2024"),
    ("Review date", "3 Sept 2026"),
    ("Declaration / signature reference", "Priyanshi Srivastava"),
):
    add_kv(t, label, value)
set_table_geometry(t)

for idx, item in enumerate(ITEMS, 1):
    if idx > 1:
        doc.add_page_break()
    doc.add_heading(f"{idx}. {item[0]}", level=1)
    t=doc.add_table(rows=0, cols=2); t.style="Table Grid"
    for label, value in zip(("Required identity","Candidate source","Source description","Identity concern","Rights","Core Nutrition","Decision choices","Decision","Reviewer initials / reference","Food-level date"), (item[0],item[1],item[2],item[3],item[4],item[5],item[6],"PENDING - reviewer must select","Priyanshi Srivastava","3 Sept 2026")):
        add_kv(t,label,value,PALE)
    set_table_geometry(t)

doc.add_heading("Source record references", level=1)
for url in ("https://fdc.nal.usda.gov/api-guide/", "https://fdc.nal.usda.gov/download-datasets/", "https://fdc.nal.usda.gov/fdc-app.html#/food-details/1750349/nutrients", "https://fdc.nal.usda.gov/fdc-app.html#/food-details/171314/nutrients", "https://fdc.nal.usda.gov/fdc-app.html#/food-details/171410/nutrients"):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); set_font(p.add_run(url), 9, False, GRAY)

doc.add_heading("Governance note", level=1)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3)
set_font(p.add_run("This task records source-identity decisions only. It does not constitute Stage A formula approval, physical-measurement approval, Stage B nutrition approval, or production release. Do not select APPROVE merely to unblock calculation. Where no exact approved source exists, CONFIRM NO MATCH is a valid outcome."), 10)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
